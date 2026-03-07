import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from models.transcript import AnalysisResult, Segment

def _format_timestamp(seconds: float) -> str:
    minutes = int(seconds) // 60
    secs = int(seconds) % 60
    return f"{minutes:02d}:{secs:02d}"


def _parse_time(time_str: str) -> float:
    parts = time_str.split(":")
    if len(parts) == 2:
        return int(parts[0]) * 60 + int(parts[1])
    elif len(parts) == 3:
        return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
    return 0.0


def _add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    new_run = paragraph._element.makeelement(qn("w:r"), {})
    rPr = paragraph._element.makeelement(qn("w:rPr"), {})
    c = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    u = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def generate_docx(
    title: str, analysis: AnalysisResult, output_dir: str,
    video_url: str = "", thumbnail_path: str | None = None,
    segments: list[Segment] | None = None,
) -> str:
    """Generate a .docx document with the analysis results."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # Title
    heading = doc.add_heading(title, level=1)

    # Video URL
    if video_url:
        url_para = doc.add_paragraph()
        _add_hyperlink(url_para, video_url, video_url)

    # Thumbnail
    if thumbnail_path and os.path.exists(thumbnail_path):
        doc.add_picture(thumbnail_path, width=Inches(5))

    # Summary
    doc.add_heading("Саммари", level=2)
    doc.add_paragraph(analysis.summary)

    # Table of contents
    doc.add_heading("Оглавление", level=2)
    for section in analysis.sections:
        time_range = f"[{section.start_time} - {section.end_time}]"
        doc.add_paragraph(
            f"{section.title} {time_range}",
            style="List Number",
        )

    # Practical action steps by section
    doc.add_heading("Пошаговые инструкции", level=2)
    for section in analysis.sections:
        if section.action_steps:
            doc.add_heading(section.title, level=3)
            for step in section.action_steps:
                doc.add_paragraph(step, style="List Bullet")

    # Translated/structured content by sections
    doc.add_heading("Транскрипт", level=2)
    for section in analysis.sections:
        time_range = f"[{section.start_time} - {section.end_time}]"
        doc.add_heading(f"{section.title} {time_range}", level=3)
        doc.add_paragraph(section.content)

    # Save
    safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in title)[:50]
    filename = f"{safe_title}.docx".strip()
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return output_path
